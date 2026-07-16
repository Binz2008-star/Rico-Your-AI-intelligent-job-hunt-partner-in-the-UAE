"""Regression tests for the adjacent-scan security hardening pass.

Covers three server-side fixes:

- FIX 1: ``client_ip_key`` must resolve the *rightmost trusted* X-Forwarded-For
  entry, not the attacker-controlled leftmost one, so per-IP rate limits cannot
  be bypassed by forging the header.
- FIX 2: the unauthenticated ``POST /chat/public`` endpoint must never let an
  unverified ``email`` adopt a real account identity — it becomes a namespaced,
  non-persisting public session, while still enforcing the registered user's
  quota gate as pure anti-dodge.
- FIX 3: ``CVParser._parse_docx`` must reject DOCX decompression bombs (a small
  zip declaring gigabytes of uncompressed content) without inflating them.

Mirrors tests/test_sse_stream_headers.py: JWT_SECRET default, module-scoped
TestClient, and a limiter-reset fixture.
"""
from __future__ import annotations

import hashlib
import io
import os
import zipfile

import pytest

os.environ.setdefault("JWT_SECRET", "ricosecret" + "x" * 21)


# ── Shared fixtures ───────────────────────────────────────────────────────────

@pytest.fixture(scope="module")
def client():
    from fastapi.testclient import TestClient
    from src.api.app import app
    return TestClient(app, raise_server_exceptions=False)


@pytest.fixture(autouse=True)
def reset_rate_limiter():
    from src.api.rate_limit import limiter
    limiter.reset()
    yield


def _make_request(xff=None, peer="9.9.9.9"):
    """Build a minimal Starlette Request with an optional X-Forwarded-For header."""
    from starlette.requests import Request

    headers = []
    if xff is not None:
        headers.append((b"x-forwarded-for", xff.encode()))
    scope = {
        "type": "http",
        "method": "GET",
        "path": "/",
        "headers": headers,
        "client": (peer, 12345),
    }
    return Request(scope)


# ── FIX 1: X-Forwarded-For rate-limit bypass ──────────────────────────────────

def test_xff_returns_rightmost_trusted_hop_not_leftmost():
    from src.api.rate_limit import client_ip_key

    # Render appends the real peer to the RIGHT; with 1 trusted hop we take the
    # rightmost entry. The leftmost ("1.1.1.1") is attacker-controlled.
    req = _make_request(xff="1.1.1.1, 2.2.2.2")
    key = client_ip_key(req)
    assert key == "2.2.2.2"
    assert key != "1.1.1.1"  # attacker-controlled leftmost must NOT win


def test_xff_honours_trusted_proxy_hops_env(monkeypatch):
    from src.api.rate_limit import client_ip_key

    # With two trusted proxies, the real client is the entry two hops from the right.
    monkeypatch.setenv("RICO_TRUSTED_PROXY_HOPS", "2")
    req = _make_request(xff="1.1.1.1, 2.2.2.2")
    assert client_ip_key(req) == "1.1.1.1"


def test_xff_fewer_entries_than_hops_falls_back_to_peer(monkeypatch):
    from src.api.rate_limit import client_ip_key

    # Misconfig/dev: header has fewer entries than expected trusted hops → do NOT
    # trust the forged head; fall back to the direct peer.
    monkeypatch.setenv("RICO_TRUSTED_PROXY_HOPS", "3")
    req = _make_request(xff="1.1.1.1, 2.2.2.2", peer="9.9.9.9")
    assert client_ip_key(req) == "9.9.9.9"


def test_no_or_empty_xff_falls_back_to_peer():
    from src.api.rate_limit import client_ip_key

    assert client_ip_key(_make_request(xff=None, peer="9.9.9.9")) == "9.9.9.9"
    assert client_ip_key(_make_request(xff="", peer="9.9.9.9")) == "9.9.9.9"
    # Only separators/whitespace → no usable entries → peer fallback.
    assert client_ip_key(_make_request(xff="  ,  ", peer="9.9.9.9")) == "9.9.9.9"


def test_trusted_proxy_hops_clamped_and_default():
    from src.api.rate_limit import _trusted_proxy_hops

    os.environ.pop("RICO_TRUSTED_PROXY_HOPS", None)
    assert _trusted_proxy_hops() == 1  # default
    os.environ["RICO_TRUSTED_PROXY_HOPS"] = "0"
    assert _trusted_proxy_hops() == 1  # clamped to >= 1
    os.environ["RICO_TRUSTED_PROXY_HOPS"] = "not-an-int"
    assert _trusted_proxy_hops() == 1  # parse error → default 1
    os.environ.pop("RICO_TRUSTED_PROXY_HOPS", None)


# ── FIX 2: unauthenticated cross-account impersonation via email ──────────────

class _FakeUser:
    def __init__(self, email):
        self.email = email


def test_public_chat_email_never_adopts_account_identity(client, monkeypatch):
    import src.repositories.users_repo as users_repo
    import src.services.subscription_gating as gating
    from src.services import chat_service

    victim_email = "victim@example.com"

    # get_user_by_email returns a registered victim account.
    monkeypatch.setattr(
        users_repo, "get_user_by_email", lambda email: _FakeUser(victim_email)
    )

    # Quota gate must still be consulted for the registered user (anti-dodge).
    gate_calls = []

    class _Gate:
        allowed = True

    def _fake_gate(user_key):
        gate_calls.append(user_key)
        return _Gate()

    monkeypatch.setattr(gating, "check_ai_message_allowed_for_user", _fake_gate)

    # Capture the session context handed to the chat service.
    captured = {}

    def _fake_send_message(ctx, message, operation_id=None, language=None):
        captured["ctx"] = ctx
        return {"message": "ok", "type": "conversational", "intent": "conversational"}

    monkeypatch.setattr(chat_service, "send_message", _fake_send_message)

    res = client.post(
        "/api/v1/rico/chat/public",
        json={"email": victim_email, "message": "hello"},
    )
    assert res.status_code == 200

    ctx = captured["ctx"]
    expected_key = "e-" + hashlib.sha256(victim_email.encode("utf-8")).hexdigest()[:40]
    # Namespaced public identity — NOT the victim's email.
    assert ctx.user_id == f"public:{expected_key}"
    assert ctx.user_id != victim_email
    assert victim_email not in ctx.user_id
    # An unverified email must never persist into a real profile.
    assert ctx.can_persist_profile is False
    assert ctx.auth_type == "public"
    # Quota gate was still enforced for the registered user (keyed on stored email).
    assert gate_calls == [victim_email]


def test_public_chat_registered_quota_gate_blocks(client, monkeypatch):
    import src.repositories.users_repo as users_repo
    import src.services.subscription_gating as gating
    from src.services import chat_service

    monkeypatch.setattr(
        users_repo, "get_user_by_email", lambda email: _FakeUser("capped@example.com")
    )

    class _BlockedGate:
        allowed = False

        def to_response(self):
            return {"message": "quota exceeded", "type": "limit", "success": False}

    monkeypatch.setattr(
        gating, "check_ai_message_allowed_for_user", lambda k: _BlockedGate()
    )

    # send_message must NOT be reached once the gate blocks.
    def _boom(*a, **k):  # pragma: no cover - must not run
        raise AssertionError("send_message called despite blocked quota gate")

    monkeypatch.setattr(chat_service, "send_message", _boom)

    res = client.post(
        "/api/v1/rico/chat/public",
        json={"email": "capped@example.com", "message": "hello"},
    )
    assert res.status_code == 200
    assert res.json().get("message") == "quota exceeded"


def test_public_chat_requires_session_or_email(client):
    res = client.post("/api/v1/rico/chat/public", json={"message": "hello"})
    assert res.status_code == 422


# ── FIX 3: DOCX decompression bomb ────────────────────────────────────────────

class _FakeInfo:
    def __init__(self, size):
        self.file_size = size


def test_parse_docx_rejects_decompression_bomb(monkeypatch):
    import src.cv_parser as cv_parser

    class _BombZip:
        def __init__(self, *a, **k):
            pass

        def infolist(self):
            # 500 MB declared uncompressed — well over the 200 MB cap.
            return [_FakeInfo(500 * 1024 * 1024)]

    monkeypatch.setattr(cv_parser.zipfile, "ZipFile", _BombZip)

    parser = cv_parser.CVParser()
    # Tiny compressed payload, huge declared inflation → rejected (empty text),
    # and crucially python-docx is never invoked to inflate it.
    assert parser._parse_docx(b"tiny-bytes") == ""


def test_parse_docx_rejects_high_ratio_bomb(monkeypatch):
    import src.cv_parser as cv_parser

    class _RatioZip:
        def __init__(self, *a, **k):
            pass

        def infolist(self):
            # 15 MB inflated from ~10 bytes → ratio far above the ceiling, and
            # total above the 10 MB floor that guards the ratio branch.
            return [_FakeInfo(15 * 1024 * 1024)]

    monkeypatch.setattr(cv_parser.zipfile, "ZipFile", _RatioZip)

    parser = cv_parser.CVParser()
    assert parser._parse_docx(b"0123456789") == ""


def test_parse_docx_accepts_normal_document():
    from docx import Document

    from src.cv_parser import CVParser

    # Build a real, small, legitimate .docx in memory.
    doc = Document()
    doc.add_paragraph("Jane Candidate")
    doc.add_paragraph("Senior Backend Engineer with 8 years experience.")
    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    # Sanity: it is a real zip and small.
    assert zipfile.is_zipfile(io.BytesIO(data))
    assert len(data) < 1024 * 1024

    text = CVParser()._parse_docx(data)
    assert "Jane Candidate" in text
    assert "Senior Backend Engineer" in text
